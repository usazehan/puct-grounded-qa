#!/usr/bin/env python3
"""Measure OCR accuracy against native source files.

Native .docx / .pdf / .xlsx from the Interchange's "Native Files (Zip)" is
ground truth; the served OCR'd PDF is what gets measured. The result licenses
two design decisions: exact numeric verification, and fuzzy span verification.

Numerics are counted per occurrence, not per unique token -- tables repeat
values, and set recall hides a dropped occurrence behind a surviving twin.
Signs are part of token identity, since (1,234) and 1,234 differ. And digit
accuracy alone is not enough: a value can survive intact while the row it
belonged to does not, which the association check catches and the guard
cannot. See check_association.

Documents failing any of these get verdict `refuse_numerics`; `--json` emits
that as a per-document flag for ingestion to carry, so numeric claims citing
them are refused by policy rather than the whole corpus losing exact matching.

Layout:
    data/raw/49421_788_1050240.PDF
    data/native/49421_788/...          (ZIP extracted into a per-item dir)

Usage:
    python scripts/ocr_accuracy.py data/raw data/native
    python scripts/ocr_accuracy.py data/raw data/native --show-errors
    python scripts/ocr_accuracy.py data/raw data/native --json data/ocr_report.json
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import warnings
from collections import Counter
from dataclasses import dataclass, field
from decimal import Decimal, InvalidOperation
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

# A corpus of agency filings contains malformed files, and their parser noise
# buries the report. Silenced by default and restored by --show-parse-errors,
# because a document that will not parse is a finding, not a nuisance: it ends
# up unmeasured, and unmeasured must not read as clean.
warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")

from puctqa.extract import extract_document  # noqa: E402
from puctqa.sources import DocumentRef  # noqa: E402

WORD_RE = re.compile(r"[A-Za-z]{3,}")

# Parenthesized negatives, currency and percent are captured, not discarded.
NUM_RE = re.compile(
    r"""
    (?P<open>\()?
    \s*\$?\s*
    (?P<sign>-)?
    (?P<num>\d[\d,]*(?:\.\d+)?)
    \s*(?P<close>\))?
    \s*(?P<pct>%)?
    """,
    re.VERBOSE,
)

# Pleading line numbers -- the 1-25 running down a testimony margin -- extract as
# standalone lines. In item 786 they are 277 of 386 numeric tokens and sit
# between every pair of content lines, so they inflate line distance and make
# small integers findable anywhere. Page furniture, not content.
PLEADING_LINE_RE = re.compile(r"\s*\d{1,2}\s*")

# Tariff prose numbers its subsections "(1) The Competitive Retailer...", and a
# parenthesized digit there is a list marker, not the accounting convention for a
# negative. Only line-initial bare integers are treated as markers: a genuine
# (1,234) adjustment sits inside a table row, never at the head of a line.
ENUMERATION_RE = re.compile(r"^[ \t]*\(\d{1,2}\)(?=\s|$)", re.MULTILINE)

# Table-of-contents lines carry section numbers and page numbers that are
# navigation, not content. Dot leaders identify them unambiguously.
TOC_LEADER_RE = re.compile(r"^.*\.{5,}.*$", re.MULTILINE)

# Below this numeric fidelity, extraction has failed systemically and the
# document should not support numeric claims at all. Above it, the misses are
# individual values, and the guard already refuses those one at a time when it
# cannot match them to a span -- excluding a 371-page tariff because six of its
# eight thousand numbers did not round-trip throws the document away to avoid
# six errors. Item 795 sits at 99.4%; a genuinely corrupted scan sits far below.
SYSTEMIC_NUMERIC_FLOOR = 99.0

# Association is different: a broken row binds a correct value to the wrong
# label, and the guard CANNOT catch it -- it verifies against the chunk, and the
# chunk is the defective artifact. So the bar stays high.
SYSTEMIC_ASSOCIATION_FLOOR = 99.5

# A bundle spreadsheet earns its way into the association check by covering the
# served document. Below this ratio of locatable pairs to served numeric tokens
# it is a workpaper behind the filing rather than the filing's own tables.
WORKPAPER_YIELD_FLOOR = 0.25

# A value this common in a document cannot be tied to one row by proximity.
MAX_VALUE_OCCURRENCES = 5

SHORT_TOKEN_MAX = 3  # short tokens carry the least redundancy against digit confusion
ASSOC_LINE_TOLERANCE = 2  # lines between a value and its row label; calibrate per corpus
ASSOC_LABEL_OVERLAP = 0.6  # fraction of label content words that must survive


# --- Numeric canonicalisation ---


def canonicalize(match: re.Match) -> str | None:
    """1,234 -> "1234"   (1,234) -> "-1234"   $1,234.50 -> "1234.5"   10.40% -> "10.4%"

    A closing paren signs the value only if an opening one was captured too, so
    "see line 5)" is not negative five.
    """
    raw = match.group("num").replace(",", "")
    try:
        value = Decimal(raw)
    except InvalidOperation:
        return None
    negative = bool(match.group("sign")) or bool(match.group("open") and match.group("close"))
    if negative:
        value = -value
    text = format(value.normalize(), "f")
    return text + ("%" if match.group("pct") else "")


def strip_furniture(text: str) -> str:
    """Remove numbering that is page apparatus rather than record content.

    Enumeration markers and TOC entries are numerals a citation would never
    need, and leaving them in lets them dominate the residual: item 795's
    apparent numeric errors were mostly "(1)" read as negative one.
    """
    text = TOC_LEADER_RE.sub("", text)
    return ENUMERATION_RE.sub("", text)


def numeric_counter(text: str) -> Counter:
    """Occurrence counts, not a set. Page apparatus excluded."""
    counts: Counter = Counter()
    for match in NUM_RE.finditer(strip_furniture(text)):
        token = canonicalize(match)
        if token is not None:
            counts[token] += 1
    return counts


def word_types(text: str) -> set[str]:
    return {w.lower() for w in WORD_RE.findall(text)}


def digits_only(token: str) -> str:
    """Digit core of a canonical token, for locating it in raw text."""
    return token.lstrip("-").rstrip("%")


# --- Ground truth readers ---


@dataclass
class NativeDoc:
    """Full text, plus (row_label, canonical token) pairs where the format has tables."""

    text: str
    pairs: list[tuple[str, str]] = field(default_factory=list)
    has_structure: bool = False
    kind: str = "unknown"


def _row_label(cells: list[str]) -> str | None:
    """Leftmost cell that reads as a label rather than a value."""
    for cell in cells:
        stripped = cell.strip()
        if len(stripped) < 3:
            continue
        if WORD_RE.search(stripped):
            return stripped
    return None


def _pairs_from_rows(rows: list[list[str]]) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = []
    for cells in rows:
        label = _row_label(cells)
        if not label:
            continue
        for cell in cells:
            if cell.strip() == label:
                continue
            for match in NUM_RE.finditer(cell):
                token = canonicalize(match)
                if token is not None:
                    pairs.append((label, token))
    return pairs


def read_docx(path: Path) -> NativeDoc:
    try:
        import docx  # python-docx
    except ImportError:
        raise SystemExit("pip install python-docx to run this script")
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    rows: list[list[str]] = []
    for table in document.tables:
        for row in table.rows:
            # Merged cells repeat per grid position in python-docx.
            seen: list[str] = []
            for cell in row.cells:
                if not seen or cell.text != seen[-1]:
                    seen.append(cell.text)
            rows.append(seen)
            parts.extend(seen)
    return NativeDoc(
        text="\n".join(parts),
        pairs=_pairs_from_rows(rows),
        has_structure=bool(rows),
        kind="docx",
    )


def read_xlsx(path: Path) -> NativeDoc:
    """Number runs and cost-of-service schedules.

    The best ground truth in the corpus: the cell grid gives row association
    exactly, so the association check stops being a proxy.
    """
    try:
        import openpyxl
    except ImportError:
        raise SystemExit("pip install openpyxl to measure spreadsheet ground truth")
    workbook = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    parts: list[str] = []
    rows: list[list[str]] = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if not any(c.strip() for c in cells):
                continue
            rows.append(cells)
            parts.append("\t".join(cells))
    workbook.close()
    return NativeDoc(
        text="\n".join(parts),
        pairs=_pairs_from_rows(rows),
        has_structure=bool(rows),
        kind="xlsx",
    )


def read_native(path: Path) -> NativeDoc:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        # Born-digital PDF: text is trustworthy, table structure is not recoverable.
        return NativeDoc(text=extract_document(path.read_bytes()).text, kind="pdf")
    if suffix in {".xlsx", ".xlsm"}:
        return read_xlsx(path)
    return read_docx(path)


NATIVE_SUFFIXES = {".docx", ".pdf", ".xlsx", ".xlsm"}

# A native ZIP holds several files, and which one is "the document" cannot be
# decided by format. Item 773 ships both a memo .docx and a memo-and-attachments
# .pdf; the served filing is the latter, so ranking .docx first would compare a
# cover memo against a document eight attachments long. That scores near 100%,
# because every word of the memo really is in the PDF -- it just measures a
# fifth of the filing and reports the fraction as the whole.
#
# So the primary is chosen by how much of the served document it accounts for,
# and format rank survives only as a tie-break. Spreadsheets are never primary:
# a cell dump is not prose, whatever it covers.
FORMAT_PREFERENCE = {".docx": 0, ".pdf": 1, ".xlsx": 2, ".xlsm": 3}
TEXT_SUFFIXES = {".docx", ".pdf"}

# A comparison needs something to compare against. Item 773's native bundle
# contains a PDF of the same scan, with no text layer -- it extracted to nothing,
# so 0 of 0 numeric tokens matched and the verdict read 100%. Absence of ground
# truth must never present as perfect agreement, so a native this thin is
# reported as no_ground_truth rather than scored.
MIN_GROUND_TRUTH_WORDS = 50
MIN_GROUND_TRUTH_NUMERICS = 20

# Absolute counts are not enough. Item 773's best native is a two-page memo:
# 191 words and 24 numeric tokens, clearing the floor above while covering 23.6%
# of a 26-page filing. It scores 100% because every figure it contains really is
# in the served text -- and the twenty-odd pages of schedules it says nothing
# about are never examined. So the native must also account for a fair share of
# the SERVED document. This is the mirror of the partial_pairing check: there the
# served file was a fragment of the native, here the native is a fragment of the
# served file.
MIN_NATIVE_COVERAGE = 60.0

# Below this word accuracy, a mispairing is likelier than bad OCR. Measured
# runs sit at 99.9-100%, so this is not a close call.
MISPAIR_FLOOR = 50.0

# ...unless the served text is almost entirely contained in the native, which
# means the pairing is right and incomplete rather than wrong. See verdict().
CONTAINMENT_FLOOR = 90.0

# Document IDs within one set run consecutively (795 set A is 1057872-1057875);
# a refiled set lands in a much later batch (1119824-1119827). A gap this large
# starts a new set. Page-range descriptions are NOT used to group -- 795 serves
# two documents both described "Pages 101 to 200".
SET_ID_GAP = 32


@dataclass
class NativeMatch:
    """Everything in an item's native bundle, sorted by what it is good for."""

    primary: Path  # the document itself: prose and numeric ground truth
    structured: list[Path] = field(default_factory=list)  # workpapers: pairs only
    candidates: list[Path] = field(default_factory=list)  # every text native in the bundle


def find_natives(native_dir: Path, ref: DocumentRef) -> NativeMatch | None:
    """All native files belonging to one item, primary first.

    Returns every match, not the first one. An item is often a bundle -- item
    786 ships four .xlsx workpapers alongside the born-digital PDF -- so a
    single-file pick either grabs a spreadsheet and scores nonsense, or grabs
    the PDF and silently discards the only files carrying table structure.
    """
    prefix = f"{ref.control_number}_{ref.item_number}_"
    matches = [
        path
        for path in native_dir.rglob("*")
        if path.suffix.lower() in NATIVE_SUFFIXES
        and not path.name.startswith("~$")
        and (
            prefix in path.name
            # ZIP contents are named descriptively, so the per-item directory is
            # the only reliable link back to an item.
            or any(prefix.rstrip("_") == part or prefix in part for part in path.parts[:-1])
        )
    ]
    if not matches:
        return None
    matches.sort(key=lambda path: (FORMAT_PREFERENCE.get(path.suffix.lower(), 9), path.name))
    return NativeMatch(
        primary=matches[0],
        structured=[m for m in matches if m.suffix.lower() in {".xlsx", ".xlsm"}],
        candidates=[m for m in matches if m.suffix.lower() in TEXT_SUFFIXES],
    )


@dataclass
class NativeScore:
    """One candidate native, scored on the axis that matters and the one that doesn't."""

    path: Path
    numeric_accuracy: float  # of ITS OWN tokens, how many the served text has
    numeric_expected: int
    word_coverage: float  # how much of the served text this native explains


def score_natives(match: NativeMatch, ocr_text: str) -> list[NativeScore]:
    """Score every text native in the bundle.

    Numeric accuracy is scoped to the native's own tokens: of the figures THIS
    document contains, how many round-trip into the served rendering. Word
    coverage is reported alongside, and is deliberately not used to choose --
    it answers a different question and choosing on it picked wrong.
    """
    ocr_words = word_types(ocr_text)
    scores: list[NativeScore] = []
    for path in match.candidates:
        try:
            text = read_native(path).text
        except Exception:
            continue
        words = word_types(text)
        if len(words) < MIN_GROUND_TRUTH_WORDS:
            continue
        result = compare_numerics(text, ocr_text)
        scores.append(
            NativeScore(
                path=path,
                numeric_accuracy=result.occurrence_accuracy,
                numeric_expected=result.occurrences_expected,
                word_coverage=(
                    len(words & ocr_words) / len(ocr_words) * 100 if ocr_words else 0.0
                ),
            )
        )
    return scores


def report_natives(match: NativeMatch, ocr_text: str) -> list[NativeScore]:
    """Every candidate above the ground-truth floors, best evidence first.

    SIX SELECTION RULES FAILED BEFORE THIS ONE STOPPED SELECTING

      - Format rank picked item 773's memo-only .docx over the
        memo-and-attachments .pdf: a fifth of the filing scoring near 100%.
      - Word coverage picked item 785's Settlement Agreement (94.7% of served
        words, 95.9% numeric) over Exhibit C (81.9%, 99.88%). Coverage optimises
        for prose; the figures live in the exhibits.
      - Unioning the bundle scored 72.9%: the ZIP holds every exhibit filed in
        the docket, while the served set is a subset of them.
      - Raw numeric agreement picked a 270-figure cover letter at 100.0% over
        Exhibit C's 4,332 of 4,337. Optimising a rate finds the native with
        least evidence.
      - Requiring a minimum figure count did not help, because 100.0 still
        outranks 99.88 before the count is ever consulted.
      - Rounding the rate so counts break near-ties would have worked for 785
        and only for 785; 99.4% and 100% still differ.

    All six tried to name THE native for a served set. Often there is not one: a
    settlement is served as an agreement plus seven exhibit workbooks, and no
    single file corresponds to the filing.

    So this reports instead of choosing. A reader sees that item 785's tariff
    exhibit round-trips 4,332 of 4,337 figures while its cover letter
    round-trips 270 of 270, and can tell which fact matters. The verdict is
    computed from the best-evidenced candidate, and every candidate is recorded
    beside it, so the judgement is visible rather than buried in a rank function.
    """
    ocr_words = word_types(ocr_text)
    scores: list[NativeScore] = []
    for path in match.candidates:
        try:
            text = read_native(path).text
        except Exception:
            continue
        words = word_types(text)
        if len(words) < MIN_GROUND_TRUTH_WORDS:
            continue
        result = compare_numerics(text, ocr_text)
        scores.append(
            NativeScore(
                path=path,
                numeric_accuracy=result.occurrence_accuracy,
                numeric_expected=result.occurrences_expected,
                word_coverage=(
                    len(words & ocr_words) / len(ocr_words) * 100 if ocr_words else 0.0
                ),
            )
        )
    # Most figures first: the verdict should rest on the candidate that puts the
    # most evidence behind its rate, not the one with the highest rate.
    scores.sort(key=lambda s: -s.numeric_expected)
    return scores


@dataclass
class ServedSet:
    """One filing, as served: usually a single PDF, sometimes several parts.

    Item 795 serves a 371-page tariff as four ~100-page PDFs, twice over. Scoring
    a part against the whole native caps its accuracy at that part's share of the
    filing -- the observed 26% was the split, not the extraction. Parts are
    concatenated in document-ID order and compared as one document.
    """

    item_number: int
    parts: list[Path]

    @property
    def label(self) -> str:
        if len(self.parts) == 1:
            return self.parts[0].name
        first = DocumentRef.from_filename(self.parts[0].name).document_id
        last = DocumentRef.from_filename(self.parts[-1].name).document_id
        return f"{self.parts[0].name.split('_')[0]}_{self.item_number} [{first}..{last}]"


def group_served_sets(paths: list[Path]) -> list[ServedSet]:
    """Group an item's served PDFs into sets by document-ID run."""
    by_item: dict[int, list[tuple[int, Path]]] = {}
    for path in paths:
        try:
            ref = DocumentRef.from_filename(path.name)
        except ValueError:
            continue
        by_item.setdefault(ref.item_number, []).append((int(ref.document_id), path))

    sets: list[ServedSet] = []
    for item, entries in sorted(by_item.items()):
        entries.sort()
        run = [entries[0]]
        for doc_id, path in entries[1:]:
            if doc_id - run[-1][0] > SET_ID_GAP:
                sets.append(ServedSet(item, [p for _, p in run]))
                run = []
            run.append((doc_id, path))
        sets.append(ServedSet(item, [p for _, p in run]))
    return sets


# --- Measurements ---


@dataclass
class NumericResult:
    occurrences_expected: int = 0
    occurrences_found: int = 0
    types_expected: int = 0
    types_found: int = 0
    short_expected: int = 0
    short_found: int = 0
    sign_losses: int = 0  # negative in native, unsigned in OCR
    missing: list[tuple[str, int]] = field(default_factory=list)

    @property
    def occurrence_accuracy(self) -> float:
        if not self.occurrences_expected:
            return 100.0
        return self.occurrences_found / self.occurrences_expected * 100

    @property
    def type_accuracy(self) -> float:
        if not self.types_expected:
            return 100.0
        return self.types_found / self.types_expected * 100

    @property
    def short_accuracy(self) -> float:
        if not self.short_expected:
            return 100.0
        return self.short_found / self.short_expected * 100


def compare_numerics(native_text: str, ocr_text: str) -> NumericResult:
    native = numeric_counter(native_text)
    ocr = numeric_counter(ocr_text)
    result = NumericResult()

    for token, expected in native.items():
        found = min(expected, ocr.get(token, 0))
        result.occurrences_expected += expected
        result.occurrences_found += found
        result.types_expected += 1
        result.types_found += 1 if ocr.get(token, 0) > 0 else 0
        if len(digits_only(token)) <= SHORT_TOKEN_MAX:
            result.short_expected += expected
            result.short_found += found
        if found < expected:
            result.missing.append((token, expected - found))
        # Surviving only without its parentheses is a sign flip, not a near miss.
        if token.startswith("-") and ocr.get(token, 0) < expected:
            unsigned = token.lstrip("-")
            if ocr.get(unsigned, 0) > native.get(unsigned, 0):
                result.sign_losses += 1

    result.missing.sort(key=lambda item: -item[1])
    return result


@dataclass
class AssociationResult:
    checked: int = 0
    intact: int = 0
    unlocatable: int = 0  # label or value absent from OCR; not scored either way
    ambiguous: int = 0  # value too common in the document to attribute by proximity
    # Distance in lines from each value to its OWN label, whether or not the pair
    # passed. This is what calibrates the tolerance -- see the histogram.
    own_distances: list[int] = field(default_factory=list)
    # Pairs where a DIFFERENT row's label is strictly closer than the correct
    # one. Raising the tolerance cannot fix these; they are real breakage.
    outranked: int = 0
    broken: list[tuple[str, str, str]] = field(default_factory=list)  # label, token, nearest

    @property
    def measurable(self) -> int:
        return self.checked - self.unlocatable - self.ambiguous

    @property
    def rate(self) -> float:
        measurable = self.measurable
        if measurable <= 0:
            return 100.0
        return self.intact / measurable * 100


def _label_lines(label: str, lines: list[str]) -> list[int]:
    """Lines where a row label is present.

    Content-word overlap rather than exact match: an rn -> m error in a label is
    a span-verification problem, not an association failure.
    """
    words = [w.lower() for w in WORD_RE.findall(label)]
    if not words:
        return []
    required = max(1, int(len(words) * ASSOC_LABEL_OVERLAP))
    return [i for i, line in enumerate(lines) if sum(w in line for w in words) >= required]


def check_association(
    pairs: list[tuple[str, str]], ocr_text: str, tolerance: int = ASSOC_LINE_TOLERANCE
) -> AssociationResult:
    """Did each (row label, value) pair survive serialization as a row?

    NEAREST PRECEDING LABEL. A value belongs to the closest row label ABOVE it
    with no other label in between -- direction, not distance.

    Three earlier models failed, each for a reason worth keeping:

      - "label within N characters" passed a column-serialized table trivially,
        because when the whole schedule fits in the window every label is near
        every value.
      - "nearest label in either direction" failed correct row-major output: a
        row's last value sits closer to the NEXT row's label than its own.
      - "within N lines" assumed a row occupies a line. Item 773 extracts one
        cell per line, currency symbols included, so a row spans a dozen lines
        and distance-to-label just measures column position. The histogram
        showed it: monotonic decline, no gap, nothing to calibrate.

    Reading order is what survives all three. Whatever the column count, the
    label comes first, so the question is which label a value falls under.
    `tolerance` is retained only as a sanity cap on how far back to look.
    """
    normalized = re.sub(r"[,$]", "", ocr_text).lower()
    lines = [l for l in normalized.split("\n") if not PLEADING_LINE_RE.fullmatch(l)]
    result = AssociationResult()

    labels = sorted({label for label, _ in pairs})

    # One label per line, and it must be the BEST match. Rate schedules share
    # vocabulary heavily -- "O&M expense" and "Depreciation expense" both match
    # any line containing "expense" -- so a first-match rule assigns rows to
    # whichever label happened to be enumerated first.
    owner_at: dict[int, str] = {}
    for i, line in enumerate(lines):
        best, best_score = None, (0.0, 0)
        for label in labels:
            words = [w.lower() for w in WORD_RE.findall(label)]
            if not words:
                continue
            hits = sum(w in line for w in words)
            if hits < max(1, int(len(words) * ASSOC_LABEL_OVERLAP)):
                continue
            # Completeness first, then specificity. Scoring by raw hit count
            # lets a long label win a line on one shared word: "Depreciation
            # expense" would take the "O&M expense" row because both contain
            # "expense" and it has more words overall.
            score = (hits / len(words), hits)
            if best is None or score > best_score:
                best, best_score = label, score
        if best:
            owner_at[i] = best

    # Locating a label and owning a line are different questions. A
    # column-serialized table puts every label on one line: all of them are
    # locatable there, but only one can govern what follows.
    label_lines = {label: _label_lines(label, lines) for label in labels}

    # Which label governs each line: the most recent one at or above it.
    governing: list[str | None] = [None] * len(lines)
    current: str | None = None
    for i in range(len(lines)):
        if i in owner_at:
            current = owner_at[i]
        governing[i] = current

    max_span = max(tolerance, 1) * 20  # sanity cap only; see docstring

    for label, token in pairs:
        needle = digits_only(token)
        result.checked += 1
        value_lines = [i for i, line in enumerate(lines) if needle and needle in line]
        own = label_lines.get(label) or []
        if not needle or not value_lines or not own:
            result.unlocatable += 1
            continue
        if len(value_lines) > MAX_VALUE_OCCURRENCES:
            result.ambiguous += 1
            continue

        distance = min(abs(v - l) for v in value_lines for l in own)
        result.own_distances.append(distance)

        intact = False
        governor = ""
        for v in value_lines:
            head = governing[v]
            if head == label:
                nearest_above = max((l for l in own if l <= v), default=None)
                if nearest_above is not None and v - nearest_above <= max_span:
                    intact = True
                    break
            governor = head or ""
        if intact:
            result.intact += 1
        else:
            result.outranked += 1
            result.broken.append((label, token, f"falls under {governor or '(no label above)'}"))
    return result


def verdict(
    num: NumericResult,
    assoc: AssociationResult,
    structured: bool,
    word_accuracy: float,
    containment: float = 100.0,
    native_words: int = MIN_GROUND_TRUTH_WORDS,
    coverage: float = 100.0,
    candidates: int = 1,
) -> str:
    """Per-document capability, in the vocabulary of the ingestion flag.

    `containment` is the mirror of word accuracy: what fraction of the SERVED
    text the native accounts for. The two together separate a wrong pairing from
    an incomplete one. Both low means two unrelated documents. Word accuracy low
    while containment is high means the served file is a genuine part of the
    native -- the pairing is right and the served side is missing parts, which
    is what 795's eight ~100-page PDFs looked like against one 371-page native.
    """
    # Nothing to compare against is not a passing grade. Checked first: every
    # accuracy below divides by a count that may be zero, and 0/0 reads as 100%.
    if (
        native_words < MIN_GROUND_TRUTH_WORDS
        or num.occurrences_expected < MIN_GROUND_TRUTH_NUMERICS
        or coverage < MIN_NATIVE_COVERAGE
    ):
        return "no_ground_truth"
    # More than one plausible native means no single number describes the set,
    # and seven rules have now tried to manufacture one. Item 785's Settlement
    # Agreement holds 12,049 figures at 95.88% because the native PDF carries
    # the agreement AND its exhibits inline; Exhibit C holds the 4,337 that were
    # actually served, at 99.89%. More figures is not better evidence when the
    # extra ones come from a superset. The candidate table carries all of it and
    # a human names the operative ground truth in the manifest, the same place
    # the 795 A/B decision lives.
    if candidates > 1:
        return "multiple_candidates"
    # A confident verdict on the wrong pair of files is the failure this whole
    # script exists to avoid, so it is checked before anything else is believed.
    if word_accuracy < MISPAIR_FLOOR:
        return "partial_pairing" if containment >= CONTAINMENT_FLOOR else "likely_mispaired"
    if num.occurrence_accuracy < SYSTEMIC_NUMERIC_FLOOR:
        return "refuse_numerics"
    # Association REPORTS, it does not gate. The rate is a lower bound: item
    # 773's residual is dominated by naming variance between the spreadsheet and
    # the printed attachment -- "Land & Land Fees" against "Land and Land Fees",
    # "Line Transformers" against "Line Transformer" -- which is label
    # morphology, not extraction failure. Separating those from real
    # misattribution ("Accumulated Depreciation" against "Total Accumulated
    # Depreciation") needs a normalization layer this script does not have.
    # Excluding a document at 100% numeric fidelity on a number that cannot
    # distinguish a plural from a broken row would be the same over-blunt
    # judgement the numeric floor already corrected.
    if not structured or not assoc.measurable:
        return "exact_unverified_structure"
    return "exact_structure_reported"


VERDICT_NOTE = {
    "multiple_candidates": (
        "the bundle holds several plausible natives and no single one corresponds to "
        "what was served -- read the candidate table and record the operative one in "
        "the manifest; numeric claims are refused by policy until someone does"
    ),
    "no_ground_truth": (
        "the paired native carries too little text to compare against -- often a PDF "
        "of the same scan, with no text layer. Nothing about this document has been "
        "verified; it is unmeasured, not clean"
    ),
    "partial_pairing": (
        "the served text is contained in the native but does not fill it -- this is "
        "part of a multi-part filing and its accuracy cannot be read on its own"
    ),
    "likely_mispaired": (
        "word accuracy too low to be OCR error -- the PDF and native file are probably "
        "not the same document; check the pairing before reading any other column"
    ),
    "exact_ok": "digits and row association intact -- exact numeric verification sound",
    "exact_unverified_structure": (
        "digits intact; native format carries no table structure, association unmeasured"
    ),
    "exact_structure_reported": (
        "digits intact; row association measured as a lower bound and reported, not "
        "gated -- see association_rate"
    ),
    "refuse_numerics": (
        "extraction failed systemically -- numeric claims citing this document should be "
        "refused by policy, not left to per-claim verification"
    ),
}


# --- Report ---


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("pdf_dir", type=Path)
    ap.add_argument("native_dir", type=Path)
    ap.add_argument("--show-errors", action="store_true")
    ap.add_argument(
        "--show-parse-errors",
        action="store_true",
        help="restore openpyxl/MuPDF warnings, silenced by default",
    )
    ap.add_argument("--json", type=Path, default=None, help="write per-document verdicts")
    ap.add_argument(
        "--line-tolerance",
        type=int,
        default=ASSOC_LINE_TOLERANCE,
        help="lines allowed between a value and its row label; calibrate per corpus",
    )
    args = ap.parse_args()
    if args.show_parse_errors:
        warnings.resetwarnings()

    served_sets = group_served_sets(sorted(args.pdf_dir.glob("*.[pP][dD][fF]")))
    matched: list[tuple[ServedSet, NativeMatch]] = []
    for served in served_sets:
        ref = DocumentRef.from_filename(served.parts[0].name)
        match = find_natives(args.native_dir, ref)
        if match:
            matched.append((served, match))

    multipart = [s for s, _ in matched if len(s.parts) > 1]
    if multipart:
        print("Multi-part filings, compared as one document each:")
        for served in multipart:
            print(f"  {served.label}  ({len(served.parts)} parts)")
        print()

    if not matched:
        print(f"No PDF/native pairs found between {args.pdf_dir} and {args.native_dir}")
        print()
        print("Native files must be linkable to an item number. Extract each ZIP")
        print("into its own directory named for the item:")
        print("    data/native/49421_786/...")
        print("    data/native/49421_788/...")
        print("Matching on control number alone is NOT done: it silently pairs the")
        print("wrong documents and produces plausible-looking but meaningless")
        print("accuracy numbers.")
        return 1

    print(
        f"{'document':<28} {'word':>6} {'num/occ':>8} {'num/set':>8} "
        f"{'short':>7} {'assoc':>7}  verdict"
    )
    print("-" * 90)

    total: Counter = Counter()
    reports: list[dict] = []
    word_errors: list[tuple[str, str]] = []
    numeric_errors: list[tuple[str, str, int]] = []
    broken_pairs: list[tuple[str, str, str, str]] = []
    anchor_summaries: list[tuple[str, dict[str, int]]] = []
    distance_histogram: list[int] = []
    unreadable: list[tuple[str, str, str]] = []

    for served, match in matched:
        parts = [extract_document(p.read_bytes()) for p in served.parts]
        ocr_text = "\n".join(part.text for part in parts)
        candidates = report_natives(match, ocr_text)
        primary_path = candidates[0].path if candidates else match.primary
        coverage = {c.path.name: round(c.word_coverage, 1) for c in candidates}
        try:
            native = read_native(primary_path)
        except Exception as exc:
            unreadable.append((served.label, primary_path.name, type(exc).__name__))
            native = NativeDoc(text="", kind="unreadable")

        # Whether a bundle spreadsheet is this document's own tables or somebody
        # else's workpaper is measurable, not assumable. Item 786's four
        # spreadsheets yielded 123,582 pairs of which 71 survived -- workpapers
        # behind a testimony. Item 773's are the model its attachments were
        # printed from. Fold in any whose values actually land in the served
        # text, and report the yield so the distinction stays visible.
        pairs = list(native.pairs)
        structured = native.has_structure
        # At most ONE workpaper. Item 773 ships two models of the same filing
        # that name identical line items differently -- "Meter Expenses" against
        # "Meter Exp", "Underground Line Expenses" against "Underground Line
        # Exp". Folding both makes line ownership a contest between synonyms:
        # one spelling wins the row and every pair carrying the other is scored
        # broken while the row is intact. Take the best-covering workbook only.
        folded: dict[str, dict] = {}
        served_tokens = numeric_counter(ocr_text)
        scored: list[tuple[float, Path, list[tuple[str, str]]]] = []
        for workpaper in match.structured:
            if workpaper == primary_path:
                continue
            try:
                workpaper_pairs = read_xlsx(workpaper).pairs
            except Exception as exc:
                # One corrupt workbook in one bundle must not lose the results
                # for every other set. Recorded, not swallowed.
                unreadable.append((served.label, workpaper.name, type(exc).__name__))
                continue
            candidate = [
                (label, token) for label, token in workpaper_pairs if token in served_tokens
            ]
            # Coverage of the served document's DISTINCT values, not pairs per
            # token -- a large model with repeated values clears a ratio of
            # counts trivially while explaining almost nothing.
            covered = {token for _, token in candidate}
            yield_rate = len(covered) / max(len(served_tokens), 1)
            if yield_rate < WORKPAPER_YIELD_FLOOR:
                continue
            scored.append((yield_rate, workpaper, candidate))

        if scored:
            yield_rate, workpaper, candidate = max(scored, key=lambda t: t[0])
            folded[workpaper.name] = {
                "pairs": len(candidate),
                "served_values_covered": round(yield_rate * 100, 1),
            }
            if len(scored) > 1:
                folded["_not_folded"] = {
                    "reason": "competing label vocabulary for the same rows",
                    "files": [w.name for r, w, _ in scored if w != workpaper],
                }
            pairs.extend(candidate)
            structured = True

        native_words, ocr_words = word_types(native.text), word_types(ocr_text)
        missing_words = native_words - ocr_words
        words_ok = len(native_words) - len(missing_words)

        word_acc = words_ok / len(native_words) * 100 if native_words else 100.0
        # What fraction of the SERVED text the native accounts for. Distinguishes
        # a wrong pairing from an incomplete one.
        containment = (
            len(native_words & ocr_words) / len(ocr_words) * 100 if ocr_words else 100.0
        )
        num = compare_numerics(native.text, ocr_text)
        assoc = check_association(pairs, ocr_text, tolerance=args.line_tolerance)
        # What fraction of the served document this native accounts for. The
        # best available native may still explain almost none of it.
        native_coverage = coverage.get(primary_path.name, 100.0)
        doc_verdict = verdict(
            num, assoc, structured, word_acc, containment,
            len(native_words), native_coverage, len(candidates),
        )

        total["words"] += len(native_words)
        total["words_ok"] += words_ok
        total["num_occ"] += num.occurrences_expected
        total["num_occ_ok"] += num.occurrences_found
        total["num_type"] += num.types_expected
        total["num_type_ok"] += num.types_found
        total["short"] += num.short_expected
        total["short_ok"] += num.short_found
        total["assoc"] += assoc.measurable
        total["assoc_ok"] += assoc.intact
        total["assoc_unlocatable"] += assoc.unlocatable + assoc.ambiguous
        total["sign_losses"] += num.sign_losses
        total["mispaired"] += doc_verdict == "likely_mispaired"
        total["partial"] += doc_verdict == "partial_pairing"
        total["outranked"] += assoc.outranked
        distance_histogram.extend(assoc.own_distances)

        word_errors.extend((served.label, w) for w in sorted(missing_words))
        numeric_errors.extend((served.label, tok, n) for tok, n in num.missing)
        broken_pairs.extend(
            (served.label, lbl, tok, near) for lbl, tok, near in assoc.broken
        )

        measurable = assoc.measurable
        assoc_display = f"{assoc.rate:>6.1f}%" if measurable else "     --"
        print(
            f"{served.label[:27]:<28} {word_acc:>5.1f}% {num.occurrence_accuracy:>7.1f}% "
            f"{num.type_accuracy:>7.1f}% {num.short_accuracy:>6.1f}% "
            f"{assoc_display}  {doc_verdict}"
        )

        reports.append(
            {
                "document": served.label,
                "parts": [p.name for p in served.parts],
                "containment": round(containment, 3),
                "native": primary_path.name,
                "native_coverage": {k: round(v, 1) for k, v in sorted(
                    coverage.items(), key=lambda kv: -kv[1]
                )},
                "native_kind": native.kind,
                "native_bundle": [w.name for w in match.structured],
                "folded_workpapers": folded,
                "word_type_accuracy": round(word_acc, 3),
                "numeric_occurrence_accuracy": round(num.occurrence_accuracy, 3),
                "numeric_type_accuracy": round(num.type_accuracy, 3),
                "numeric_occurrences": num.occurrences_expected,
                "short_token_accuracy": round(num.short_accuracy, 3),
                "short_tokens": num.short_expected,
                "sign_losses": num.sign_losses,
                "association_rate": round(assoc.rate, 3) if measurable else None,
                "association_pairs_measured": measurable,
                "association_pairs_unlocatable": assoc.unlocatable,
                "association_pairs_ambiguous": assoc.ambiguous,
                "verdict": doc_verdict,
                "verdict_note": VERDICT_NOTE[doc_verdict],
                "numeric_verifiable": doc_verdict.startswith("exact_"),
                "native_coverage_pct": round(native_coverage, 1),
                "verified_through": primary_path.name,
                # Every candidate, not just the one the verdict rests on. Item
                # 785's tariff exhibit and its cover letter tell different
                # stories and both belong in the record.
                "candidates": [
                    {
                        "native": c.path.name,
                        "numeric_accuracy": round(c.numeric_accuracy, 3),
                        "numeric_expected": c.numeric_expected,
                        "word_coverage": round(c.word_coverage, 1),
                    }
                    for c in candidates
                ],
                "native_words": len(native_words),
                "native_numerics": num.occurrences_expected,
            }
        )
        merged: Counter = Counter()
        for part in parts:
            merged.update(part.anchor_coverage())
        anchor_summaries.append((served.label, dict(merged)))

    def pct(ok: str, n: str) -> float:
        return total[ok] / total[n] * 100 if total[n] else 100.0

    print("=" * 90)
    if total["mispaired"]:
        print(
            f"STOP: {total['mispaired']} document(s) look mispaired with their native file."
        )
        print("Every figure below is aggregated across them and is meaningless until the")
        print("pairing is fixed. Re-check that each ZIP was extracted into a directory")
        print("named for its own item.")
        print()
    if total["partial"]:
        print(
            f"NOTE: {total['partial']} filing(s) are contained in their native but do not"
        )
        print("fill it -- served parts are missing. Their accuracy is not readable.")
        print()
    print(f"Documents compared:        {len(matched)}")
    print(
        f"Word-type accuracy:        {pct('words_ok','words'):.2f}%  "
        f"({total['words_ok']}/{total['words']})"
    )
    print(
        f"Numeric, per occurrence:   {pct('num_occ_ok','num_occ'):.2f}%  "
        f"({total['num_occ_ok']}/{total['num_occ']})  <- licenses exact matching"
    )
    print(
        f"Numeric, unique types:     {pct('num_type_ok','num_type'):.2f}%  "
        f"({total['num_type_ok']}/{total['num_type']})  <- what the old script reported"
    )
    gap = pct("num_type_ok", "num_type") - pct("num_occ_ok", "num_occ")
    if gap > 0.05:
        print(f"  Gap of {gap:.2f} points: repeated values are masking per-occurrence errors.")
        print("  The wider this gap, the more table-like the document, and the less the")
        print("  set-based number means.")
    print(
        f"Short tokens (<={SHORT_TOKEN_MAX} chars):  {pct('short_ok','short'):.2f}%  "
        f"({total['short_ok']}/{total['short']})  <- excluded entirely by the old filter"
    )
    if total["assoc"]:
        print(
            f"Row association intact:    {pct('assoc_ok','assoc'):.2f}%  "
            f"({total['assoc_ok']}/{total['assoc']})  <- value falls under its own row label"
        )
        if total["assoc_unlocatable"]:
            print(
                f"  ({total['assoc_unlocatable']} pair(s) unmeasurable: label or value "
                "not locatable, or value too common to attribute)"
            )
    else:
        print("Row association:           not measurable (no structured native files paired)")
    if distance_histogram:
        print()
        print("Line distance from a value to its own row label:")
        buckets = [(0, 0), (1, 1), (2, 2), (3, 3), (4, 5), (6, 9), (10, 19), (20, 10**9)]
        width = max(sum(1 for d in distance_histogram if lo <= d <= hi) for lo, hi in buckets)
        for lo, hi in buckets:
            n = sum(1 for d in distance_histogram if lo <= d <= hi)
            name = f"{lo}" if lo == hi else (f"{lo}+" if hi > 10**8 else f"{lo}-{hi}")
            bar = "#" * int(n / max(width, 1) * 40)
            print(f"  {name:>6} {n:>7}  {bar}")
        print("  Diagnostic only: association is decided by reading order, not distance.")
        print("  The absence of a gap here is why -- distance from a label just measures")
        print("  column position when a row extracts one cell per line.")
    if total["sign_losses"]:
        print(
            f"Sign losses:               {total['sign_losses']}  "
            "(parenthesized negative read as positive)"
        )
    print()

    # Judged per document, never pooled. A corpus rate mixes documents with
    # different provenance -- 773's clean tables with 795-B's delta against a
    # sibling's native -- and describes none of them. Same reason the split
    # parts of 795 could not be read individually.
    multi = [r for r in reports if len(r.get("candidates", [])) > 1]
    if multi:
        print(f"Sets with several plausible natives ({len(multi)}) -- verdict deferred:")
        for r in multi:
            print(f"  {r['document'][:44]}")
            for c in r["candidates"][:6]:
                print(
                    f"     {c['numeric_accuracy']:>6.2f}% of {c['numeric_expected']:>6} "
                    f"figures   {c['word_coverage']:>5.1f}% of text   {c['native'][:44]}"
                )
        print("  No single file is 'the' native for a set served as an agreement plus")
        print("  its exhibits, and more figures is not better evidence when the extra")
        print("  ones come from a superset. Name the operative native in the manifest;")
        print("  until then these sets refuse numeric claims by policy.")
        print()
    if unreadable:
        print(f"Native files that could not be parsed ({len(unreadable)}):")
        for doc, name, kind in unreadable:
            print(f"  {doc[:30]:<32} {name[:44]:<46} {kind}")
        print("Those files contributed nothing. If one was a set's only ground truth,")
        print("the set is unmeasured -- check its verdict below rather than assuming.")
        print()
    unmeasured = [r["document"] for r in reports if r["verdict"] == "no_ground_truth"]
    if unmeasured:
        print("STOP: no usable ground truth for:")
        for name in unmeasured:
            print(f"  {name}")
        print("No native in the bundle accounts for enough of the served document to")
        print("verify it -- either it carries almost no text (a PDF of the same scan) or")
        print("it is a short memo beside a long filing. These documents are UNMEASURED,")
        print("not clean: a 100% here means every figure in a fragment matched, while")
        print("the pages that fragment says nothing about were never examined.")
        print()
    failing = [r["document"] for r in reports if r["verdict"] == "refuse_numerics"]
    numeric = pct("num_occ_ok", "num_occ")
    if not failing:
        print("Digits survive extraction -> exact numeric verification is sound on")
        print("every document compared.")
        if numeric < 100.0:
            print(
                f"  The {total['num_occ'] - total['num_occ_ok']} miss(es) are individual "
                "values, refused per claim by the guard,"
            )
            print("  not grounds for excluding a document. That residual is the expected")
            print("  refusal rate from extraction causes.")
    else:
        print(
            f"WARNING: {len(failing)} of {len(reports)} document(s) fell below "
            f"{SYSTEMIC_NUMERIC_FLOOR}% numeric fidelity."
        )
        for name in failing:
            print(f"  {name}")
        print("Refuse numeric claims citing those. The other documents are unaffected --")
        print("the corpus-wide rate above pools them and describes neither.")

    if total["assoc"] and pct("assoc_ok", "assoc") < 100.0:
        print()
        print("Row association is a LOWER BOUND, reported and not gated. Its residual")
        print("mixes real misattribution -- a value under a subtotal instead of its line,")
        print("which the guard cannot catch -- with naming variance between the native")
        print("and the served rendering ('Land & Land Fees' / 'Land and Land Fees').")
        print("Separating them needs label normalization this script does not do.")

    if pct("words_ok", "words") < 100.0:
        print("Prose differs from the native -> span verification must be fuzzy, not exact.")
        print("  (On a text-layer corpus this is lost whitespace and native typos, not OCR.)")

    if anchor_summaries:
        print()
        print("Citation anchor coverage (pages by scheme):")
        for name, counts in anchor_summaries:
            print(
                f"  {name[:30]:<31} bates={counts['bates']:>3} "
                f"label={counts['page_label']:>3} pdf_page={counts['pdf_page']:>3}"
            )
        weak = sum(c["pdf_page"] for _, c in anchor_summaries)
        if weak:
            print(
                f"\n  {weak} page(s) fall back to PDF page numbering. Citations to"
                "\n  those pages must be flagged as weakly anchored."
            )

    if args.show_errors:
        if numeric_errors:
            print()
            print(f"Numeric tokens in the native but absent from the served text "
                  f"({len(numeric_errors)}):")
            for doc, tok, n in numeric_errors[:40]:
                print(f"  {doc[:30]:<31} {tok:>16}  x{n}")
        if broken_pairs:
            print()
            print(f"Row associations broken ({len(broken_pairs)}):")
            print(f"  {'document':<25} {'native row label':<28} {'value':>12}   nearest label in OCR")
            for doc, label, tok, near in broken_pairs[:40]:
                print(f"  {doc[:24]:<25} {label[:27]:<28} {tok:>12}   {near[:44]}")
        if word_errors:
            print()
            print(f"Words in native text but absent from OCR ({len(word_errors)}):")
            for doc, word in word_errors[:40]:
                print(f"  {doc[:30]:<31} {word}")

    if args.json:
        args.json.parent.mkdir(parents=True, exist_ok=True)
        args.json.write_text(json.dumps({"documents": reports}, indent=2))
        print()
        print(f"Wrote per-document verdicts to {args.json}")
        print("Ingestion should carry `numeric_verifiable` onto each document row and")
        print("refuse numeric claims cited to documents where it is false.")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())